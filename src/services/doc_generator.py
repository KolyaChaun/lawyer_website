import datetime
from pathlib import Path

from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Cm, Pt

MEDIA_DIR = Path("media/temp")
MEDIA_DIR.mkdir(parents=True, exist_ok=True)


def generate_advocate_request_doc(form_data: dict) -> str:
    timestamp = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
    safe_name = form_data.get("fullName", "user").replace(" ", "_")
    filename = MEDIA_DIR / f"{timestamp}_{safe_name}.docx"

    doc = Document()

    for section in doc.sections:
        section.top_margin = Cm(1.5)
        section.bottom_margin = Cm(2)

    style = doc.styles["Normal"]
    font = style.font
    font.name = "Times New Roman"
    font.size = Pt(12)

    lines = [
        form_data.get("organizationName", "-"),
        f"Код ЄДРПОУ- {form_data.get('organizationEDRPOU', '-')}",
        f"Поштова адреса: {form_data.get('organizationAddress', '-')}",
        f"Телефон: {form_data.get('organizationContact', '-')}",
        "",
        "Адвокат Рудик Марк Романович",
        "свідоцтво про право на зайняття адвокатською",
        "діяльністю № 001971 від 06.05.2022, видане",
        "Радою адвокатів Миколаївської області",
        "номер телефону – 0632258425",
        "e-mail: rudikmark01@gmail.com",
        "адреса для листування: 65016, м. Одеса, вул. Баштанна, 2",
        "",
        "Діє в інтересах:",
        form_data.get("fullName", "-"),
        f"Поштова адреса: {form_data.get('address', '-')}",
        f"Телефон: {form_data.get('phone', '-')}",
        f"Електронна адреса: {form_data.get('email', '-')}",
        "",
    ]

    bold_lines = [
        form_data.get("organizationName", "-"),
        "Адвокат Рудик Марк Романович",
        form_data.get("fullName", "-"),
    ]
    for line in lines:
        p = doc.add_paragraph()
        run = p.add_run(line)
        if line in bold_lines:
            run.bold = True
        p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        p.paragraph_format.left_indent = Cm(8.5)
        p.paragraph_format.right_indent = Cm(-2)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1

    # ===== Заголовок =====

    title_p = doc.add_paragraph()
    title_run = title_p.add_run("Адвокатський запит")
    title_run.bold = True
    title_run.font.size = Pt(12)
    title_p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    title_p.paragraph_format.space_after = Pt(12)

    # p1
    p1 = doc.add_paragraph(
        "Відповідно до п. 2. ст. 24 Закону України «Про адвокатуру та адвокатську діяльність» "
        "від 05.07.2012 - орган державної влади, орган місцевого самоврядування, їх посадові "
        "та службові особи, керівники підприємств, установ, організацій, громадських об’єднань, "
        "яким направлено адвокатський запит, зобов’язані не пізніше п’яти робочих днів з дня "
        "отримання запиту надати адвокату відповідну інформацію, копії документів."
    )
    p1.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    p1.paragraph_format.left_indent = Cm(-1)
    p1.paragraph_format.right_indent = Cm(-1)
    p1.paragraph_format.first_line_indent = Cm(1.25)
    p1.paragraph_format.space_after = Pt(8)
    p1.paragraph_format.line_spacing = 1

    # p2
    p2 = doc.add_paragraph(
        "Відповідно до п. 3. ст. 24 Закону України «Про адвокатуру та адвокатську діяльність» "
        "від 05.07.2012 - відмова в наданні інформації на адвокатський запит, несвоєчасне або "
        "неповне надання інформації, надання інформації, що не відповідає дійсності, тягнуть "
        "за собою відповідальність, встановлену законом."
    )
    p2.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    p2.paragraph_format.left_indent = Cm(-1)
    p2.paragraph_format.right_indent = Cm(-1)
    p2.paragraph_format.first_line_indent = Cm(1.25)
    p2.paragraph_format.space_after = Pt(8)
    p2.paragraph_format.line_spacing = 1

    # p3
    current_date = datetime.datetime.now().strftime("%d.%m.%Y")
    p3 = doc.add_paragraph()

    p3.add_run("Я, адвокат Рудик Марк Романович, представляю інтереси ")

    full_name_run = p3.add_run(f"{form_data.get('fullName', '____________')} ")
    full_name_run.bold = True

    p3.add_run(
        f"на підставі договору про надання правової допомоги від {current_date}."
    )

    p3.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    p3.paragraph_format.left_indent = Cm(-1)
    p3.paragraph_format.right_indent = Cm(-1)
    p3.paragraph_format.first_line_indent = Cm(1.25)
    p3.paragraph_format.space_after = Pt(12)
    p3.paragraph_format.line_spacing = 1

    all_paragraphs = [p1, p2, p3]

    claim_text = form_data.get("claimText", "").strip()
    if claim_text:
        spacer_top = doc.add_paragraph()
        claim_p = doc.add_paragraph(claim_text)
        spacer_bottom = doc.add_paragraph()
        all_paragraphs.extend([spacer_top, claim_p, spacer_bottom])

    # p4
    p4 = doc.add_paragraph(
        "З огляду на викладене, з метою здійснення належного захисту прав клієнта в органах державної влади, "
        "а також у судах усіх інстанцій, керуючись Законом України «Про адвокатуру та адвокатську діяльність» "
        "від 05.07.2012, Законом України «Про інформацію», Конституцією України,-"
    )
    all_paragraphs.append(p4)

    proshu_p = doc.add_paragraph()
    proshu_run = proshu_p.add_run("ПРОШУ:")
    proshu_run.bold = True
    proshu_run.font.size = Pt(12)
    proshu_p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    proshu_p.paragraph_format.space_after = Pt(12)
    all_paragraphs.append(proshu_p)

    documents_requested = form_data.get("documentsRequested", "").strip()
    if documents_requested:
        p_item1 = doc.add_paragraph()
        run_num1 = p_item1.add_run("1. ")
        run_num1.bold = True
        run_text1 = p_item1.add_run(
            f"Надати інформацію та відповідні підтверджуючі документи: {documents_requested}."
        )
        p_item1.paragraph_format.first_line_indent = Cm(0.5)
        p_item1.paragraph_format.space_before = Pt(0)
        p_item1.paragraph_format.space_after = Pt(0)
        p_item1.paragraph_format.line_spacing = 1
        all_paragraphs.append(p_item1)

        p_item2 = doc.add_paragraph()
        run_num2 = p_item2.add_run("2. ")
        run_num2.bold = True
        run_text2 = p_item2.add_run(
            "Відповідь на адвокатський запит, копії документів та інформацію прошу надіслати на мою електронну пошту: rudikmark01@gmail.com."
        )
        p_item2.paragraph_format.first_line_indent = Cm(0.5)
        p_item2.paragraph_format.space_before = Pt(0)
        p_item2.paragraph_format.space_after = Pt(0)
        p_item2.paragraph_format.line_spacing = 1
        all_paragraphs.append(p_item2)

    # format
    for p in all_paragraphs:
        p.paragraph_format.line_spacing = 1
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.left_indent = Cm(-1)
        p.paragraph_format.right_indent = Cm(-1)
        p.paragraph_format.first_line_indent = Cm(1)

    attachments_title = doc.add_paragraph()
    attachments_run = attachments_title.add_run("Додатки:")
    attachments_run.bold = True
    attachments_title.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    attachments_title.paragraph_format.space_before = Pt(12)
    attachments_title.paragraph_format.space_after = Pt(6)
    attachments_title.paragraph_format.line_spacing = 1
    all_paragraphs.append(attachments_title)

    p_attach1 = doc.add_paragraph("1. Копія ордеру.")
    p_attach1.paragraph_format.first_line_indent = Cm(0.5)
    p_attach1.paragraph_format.space_before = Pt(0)
    p_attach1.paragraph_format.space_after = Pt(0)
    p_attach1.paragraph_format.line_spacing = 1
    all_paragraphs.append(p_attach1)

    p_attach2 = doc.add_paragraph(
        "2. Копія свідоцтва на право зайняття адвокатською діяльністю."
    )
    p_attach2.paragraph_format.first_line_indent = Cm(0.5)
    p_attach2.paragraph_format.space_before = Pt(0)
    p_attach2.paragraph_format.space_after = Pt(0)
    p_attach2.paragraph_format.line_spacing = 1
    all_paragraphs.append(p_attach2)

    doc.add_paragraph().paragraph_format.space_before = Pt(12)

    current_date = datetime.datetime.now().strftime("%d.%m.%Y")

    table = doc.add_table(rows=1, cols=2)
    table.autofit = True

    cell_left = table.cell(0, 0)
    cell_left.text = current_date
    cell_left.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    cell_left.paragraphs[0].runs[0].bold = True

    cell_right = table.cell(0, 1)
    cell_right.text = "адвокат Марк РУДИК"
    cell_right.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    cell_right.paragraphs[0].runs[0].bold = True

    doc.save(filename)
    return str(filename)
